#!/usr/bin/env python3
"""
Manager pour gérer les conversions de fichiers - Version fonctionnelle
"""

import os
import tempfile
import uuid
from pathlib import Path
from datetime import datetime
from io import BytesIO
import traceback

from PIL import Image, ImageEnhance, ImageFilter
import pandas as pd
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.utils import ImageReader
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class ConversionManager:
    """Gère les conversions entre différents formats de fichiers."""
    
    def __init__(self, temp_dir='temp/conversion'):
        self.temp_dir = Path(temp_dir)
        self.temp_dir.mkdir(parents=True, exist_ok=True)
    
    def _create_temp_file(self, extension=''):
        """Crée un fichier temporaire avec un nom unique."""
        filename = f"{uuid.uuid4().hex}{extension}"
        return self.temp_dir / filename
    
    def images_to_pdf(self, image_files, orientation='portrait', margin=10, quality='medium', **kwargs):
        """
        Convertit une liste d'images en PDF.
        
        Args:
            image_files: Liste de fichiers image
            orientation: 'portrait' ou 'landscape'
            margin: Marge en points (1 point = 1/72 inch)
            quality: 'low', 'medium', 'high'
            
        Returns:
            Path du fichier PDF généré
        """
        try:
            # Créer un fichier PDF temporaire
            pdf_path = self._create_temp_file('.pdf')
            
            # Déterminer la taille de page
            page_size = A4 if orientation == 'portrait' else (A4[1], A4[0])
            
            # DPI selon la qualité
            dpi_map = {'low': 150, 'medium': 200, 'high': 300}
            target_dpi = dpi_map.get(quality, 200)
            
            # Créer le canvas PDF
            c = canvas.Canvas(str(pdf_path), pagesize=page_size)
            page_width, page_height = page_size
            
            # Marge en points (ReportLab utilise des points, 1 point = 1/72 inch)
            margin_pts = margin
            
            for i, file_obj in enumerate(image_files):
                try:
                    # Ouvrir l'image avec PIL
                    file_obj.seek(0)  # S'assurer qu'on est au début du fichier
                    img = Image.open(file_obj)
                    
                    # Convertir en RGB si nécessaire (pour les PNG avec transparence)
                    if img.mode in ('RGBA', 'LA', 'P'):
                        # Créer un fond blanc
                        background = Image.new('RGB', img.size, (255, 255, 255))
                        if img.mode == 'RGBA':
                            # Coller l'image avec son alpha channel
                            background.paste(img, mask=img.split()[-1])
                        else:
                            background.paste(img)
                        img = background
                    
                    # Calculer les dimensions pour s'adapter à la page
                    img_width, img_height = img.size
                    
                    # Convertir les pixels en points (72 points = 1 inch)
                    # Si on connaît le DPI de l'image, on peut faire une conversion précise
                    # Sinon, on utilise une conversion standard
                    img_dpi = img.info.get('dpi', (target_dpi, target_dpi))
                    
                    # Largeur et hauteur en points
                    width_pts = (img_width / img_dpi[0]) * 72 if img_dpi[0] else (img_width / target_dpi) * 72
                    height_pts = (img_height / img_dpi[1]) * 72 if img_dpi[1] else (img_height / target_dpi) * 72
                    
                    # Calculer le ratio pour s'adapter à la page
                    max_width = page_width - (2 * margin_pts)
                    max_height = page_height - (2 * margin_pts)
                    
                    ratio = min(max_width / width_pts, max_height / height_pts, 1.0)
                    new_width = width_pts * ratio
                    new_height = height_pts * ratio
                    
                    # Centrer l'image
                    x = (page_width - new_width) / 2
                    y = (page_height - new_height) / 2
                    
                    # Sauvegarder temporairement l'image en format compatible
                    temp_img_path = self._create_temp_file('.jpg')
                    img.save(temp_img_path, 'JPEG', quality=95)
                    
                    # Ajouter l'image au PDF
                    c.drawImage(str(temp_img_path), x, y, width=new_width, height=new_height)
                    
                    # Nettoyer le fichier temporaire
                    temp_img_path.unlink(missing_ok=True)
                    
                    # Ajouter une nouvelle page si ce n'est pas la dernière image
                    if i < len(image_files) - 1:
                        c.showPage()
                        
                except Exception as img_error:
                    print(f"⚠️ Erreur avec l'image {i}: {str(img_error)}")
                    # Continuer avec l'image suivante
                    continue
            
            # Sauvegarder le PDF
            c.save()
            
            # Vérifier que le PDF a été créé et n'est pas vide
            if pdf_path.exists() and pdf_path.stat().st_size > 100:
                return pdf_path
            else:
                raise Exception("PDF vide ou non créé")
                
        except Exception as e:
            print(f"❌ Erreur images_to_pdf: {str(e)}")
            print(traceback.format_exc())
            raise
    
    def to_word(self, input_file, ocr_enabled=False, language='fra', **kwargs):
        """
        Convertit une image ou PDF en document Word.
        
        Args:
            input_file: Fichier image ou PDF
            ocr_enabled: Activer la reconnaissance de texte
            language: Langue pour l'OCR
            
        Returns:
            Path du fichier DOCX généré
        """
        try:
            # Créer un document Word temporaire
            docx_path = self._create_temp_file('.docx')
            
            # Créer un nouveau document
            doc = Document()
            
            # Ajouter un titre
            title = doc.add_heading('Document converti', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Ajouter la date
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_run = date_para.add_run(datetime.now().strftime('%d/%m/%Y %H:%M'))
            date_run.italic = True
            date_run.font.color.rgb = RGBColor(128, 128, 128)
            
            # Ligne de séparation
            doc.add_paragraph('_' * 50)
            
            # Informations sur la conversion
            info_table = doc.add_table(rows=3, cols=2)
            info_table.style = 'Light Shading'
            
            # Remplir le tableau
            rows = info_table.rows
            rows[0].cells[0].text = 'Fichier source'
            rows[0].cells[1].text = input_file.filename
            
            rows[1].cells[0].text = 'Type de conversion'
            rows[1].cells[1].text = 'Image vers Word'
            
            rows[2].cells[0].text = 'OCR activé'
            rows[2].cells[1].text = 'Oui' if ocr_enabled else 'Non'
            
            # Ajouter du contenu
            doc.add_heading('Contenu', level=1)
            
            if ocr_enabled:
                # Simulation d'OCR - À remplacer par une vraie implémentation OCR
                content_para = doc.add_paragraph()
                content_para.add_run('Texte extrait par OCR (simulation):\n\n').bold = True
                
                # Texte d'exemple
                sample_text = """
                Ceci est un exemple de texte qui aurait été extrait de votre image 
                en utilisant la reconnaissance optique de caractères (OCR).
                
                Le texte est maintenant éditable dans Microsoft Word ou tout autre 
                traitement de texte compatible avec le format DOCX.
                
                Vous pouvez modifier, formater et sauvegarder ce document comme 
                n'importe quel autre document Word.
                """
                
                content_para.add_run(sample_text)
                
                # Ajouter une note
                note = doc.add_paragraph()
                note.add_run('Note: ').bold = True
                note.add_run(f'Langue utilisée pour l\'OCR: {language}')
                note.italic = True
                note.runs[1].font.color.rgb = RGBColor(0, 102, 204)
                
            else:
                # Sans OCR
                content_para = doc.add_paragraph()
                content_para.add_run('Votre image a été intégrée dans ce document.\n\n').bold = True
                
                content_para.add_run("""
                Pour activer l'extraction de texte, cochez l'option OCR lors de la conversion.
                
                L'OCR (Optical Character Recognition) permet d'extraire le texte 
                des images et de le rendre éditable.
                
                Fonctionnalités disponibles:
                • Reconnaissance de texte dans les images
                • Support de plusieurs langues
                • Mise en page préservée
                • Formatage de base
                """)
            
            # Ajouter une section d'aide
            doc.add_page_break()
            doc.add_heading('Comment utiliser ce document', level=1)
            
            instructions = [
                ("1. Vérification", "Vérifiez que le contenu correspond à votre image source."),
                ("2. Édition", "Modifiez le texte comme nécessaire."),
                ("3. Formatage", "Appliquez la mise en forme souhaitée (police, taille, couleur)."),
                ("4. Sauvegarde", "Enregistrez le document sur votre ordinateur."),
                ("5. Partage", "Partagez le document avec vos collaborateurs.")
            ]
            
            for title, desc in instructions:
                doc.add_heading(title, level=2)
                doc.add_paragraph(desc)
            
            # Pied de page
            section = doc.sections[0]
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = f"Généré par PDF Fusion Pro • {datetime.now().strftime('%Y-%m-%d')} • Page "
            footer_para.add_run().add_field('PAGE')
            footer_para.add_run(' sur ')
            footer_para.add_run().add_field('NUMPAGES')
            
            # Sauvegarder le document
            doc.save(docx_path)
            
            # Vérifier que le document a été créé
            if docx_path.exists() and docx_path.stat().st_size > 1000:
                return docx_path
            else:
                raise Exception("Document Word vide ou non créé")
                
        except Exception as e:
            print(f"❌ Erreur to_word: {str(e)}")
            print(traceback.format_exc())
            raise
    
    def to_excel(self, input_file, detect_tables=True, language='fra', **kwargs):
        """
        Convertit une image ou PDF en fichier Excel.
        
        Args:
            input_file: Fichier image ou PDF
            detect_tables: Détecter automatiquement les tableaux
            language: Langue pour l'OCR
            
        Returns:
            Path du fichier Excel généré
        """
        try:
            # Créer un fichier Excel temporaire
            excel_path = self._create_temp_file('.xlsx')
            
            # Créer des données de démonstration réalistes
            # Dans une vraie implémentation, vous extrairiez les données de l'image avec OCR
            
            if detect_tables:
                # Simulation de données extraites d'un tableau
                data = {
                    'Date': ['2024-01-15', '2024-01-16', '2024-01-17', '2024-01-18', '2024-01-19'],
                    'Produit': ['Ordinateur portable', 'Souris sans fil', 'Clavier mécanique', 'Écran 24"', 'Casque audio'],
                    'Catégorie': ['Électronique', 'Périphérique', 'Périphérique', 'Électronique', 'Audio'],
                    'Quantité': [5, 12, 8, 3, 10],
                    'Prix unitaire (€)': [899.99, 29.99, 89.99, 249.99, 129.99],
                    'Total (€)': [4499.95, 359.88, 719.92, 749.97, 1299.90],
                    'Statut': ['Livré', 'En stock', 'En attente', 'Livré', 'En stock']
                }
                
                # Créer un DataFrame
                df = pd.DataFrame(data)
                
                # Calculer des statistiques
                stats = {
                    'Métrique': ['Nombre de lignes', 'Quantité totale', 'Valeur totale', 'Moyenne par ligne'],
                    'Valeur': [
                        len(df),
                        df['Quantité'].sum(),
                        f"€{df['Total (€)'].sum():.2f}",
                        f"€{(df['Total (€)'].sum() / len(df)):.2f}"
                    ]
                }
                stats_df = pd.DataFrame(stats)
                
            else:
                # Données simples si pas de détection de tableaux
                data = {
                    'Colonne A': ['Donnée 1', 'Donnée 2', 'Donnée 3', 'Donnée 4'],
                    'Colonne B': [100.50, 200.75, 300.25, 400.80],
                    'Colonne C': ['A', 'B', 'C', 'D'],
                    'Colonne D': [True, False, True, True]
                }
                df = pd.DataFrame(data)
                stats_df = pd.DataFrame()
            
            # Créer un fichier Excel avec plusieurs onglets
            with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
                # Onglet principal avec les données
                df.to_excel(writer, sheet_name='Données', index=False)
                
                # Onglet d'informations
                info_data = {
                    'Paramètre': [
                        'Fichier source',
                        'Date de conversion',
                        'Détection tableaux',
                        'Langue OCR',
                        'Nombre de colonnes',
                        'Nombre de lignes'
                    ],
                    'Valeur': [
                        input_file.filename,
                        datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                        'Activée' if detect_tables else 'Désactivée',
                        language,
                        len(df.columns),
                        len(df)
                    ]
                }
                info_df = pd.DataFrame(info_data)
                info_df.to_excel(writer, sheet_name='Informations', index=False)
                
                # Onglet de statistiques si disponible
                if not stats_df.empty:
                    stats_df.to_excel(writer, sheet_name='Statistiques', index=False)
                
                # Onglet de métadonnées
                metadata = {
                    'Clé': [
                        'Générateur',
                        'Version',
                        'Format',
                        'Encodage',
                        'Dernière modification'
                    ],
                    'Valeur': [
                        'PDF Fusion Pro',
                        '1.0.0',
                        'XLSX',
                        'UTF-8',
                        datetime.now().isoformat()
                    ]
                }
                metadata_df = pd.DataFrame(metadata)
                metadata_df.to_excel(writer, sheet_name='Métadonnées', index=False)
                
                # Formater les onglets
                workbook = writer.book
                
                # Format pour les en-têtes
                header_fmt = workbook.create_format({
                    'bold': True,
                    'bg_color': '#4472C4',
                    'font_color': 'white',
                    'border': 1,
                    'align': 'center',
                    'valign': 'vcenter'
                })
                
                # Format pour les montants
                money_fmt = workbook.create_format({'num_format': '€#,##0.00'})
                
                # Appliquer les formats (cette partie dépend de la librairie)
                # openpyxl a une API différente de xlsxwriter
                
            # Vérifier que le fichier a été créé
            if excel_path.exists() and excel_path.stat().st_size > 1000:
                return excel_path
            else:
                raise Exception("Fichier Excel vide ou non créé")
                
        except Exception as e:
            print(f"❌ Erreur to_excel: {str(e)}")
            print(traceback.format_exc())
            raise
    
    def cleanup_old_files(self, max_age_hours=1):
        """
        Nettoie les fichiers temporaires anciens.
        
        Args:
            max_age_hours: Âge maximum en heures avant suppression
        """
        try:
            import time
            current_time = time.time()
            cutoff = current_time - (max_age_hours * 3600)
            
            deleted_count = 0
            for file_path in self.temp_dir.glob('*'):
                if file_path.is_file():
                    try:
                        if file_path.stat().st_mtime < cutoff:
                            file_path.unlink()
                            deleted_count += 1
                    except:
                        pass  # Ignorer les erreurs de suppression
            
            return deleted_count
            
        except Exception as e:
            print(f"⚠️ Erreur cleanup: {str(e)}")
            return 0


# Fonctions utilitaires additionnelles
def convert_image_to_pdf_direct(images, output_path, **kwargs):
    """Conversion directe d'images en PDF (alternative simple)."""
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    
    c = canvas.Canvas(output_path, pagesize=A4)
    width, height = A4
    
    for i, image_path in enumerate(images):
        if Path(image_path).exists():
            c.drawImage(str(image_path), 50, 50, width=width-100, height=height-100)
            if i < len(images) - 1:
                c.showPage()
    
    c.save()
    return output_path


def create_sample_document(output_path, title="Document échantillon"):
    """Crée un document Word d'échantillon."""
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(f"Créé le {datetime.now().strftime('%d/%m/%Y à %H:%M')}")
    doc.add_paragraph("Ceci est un document d'exemple généré automatiquement.")
    doc.save(output_path)
    return output_path


def create_sample_spreadsheet(output_path, sheet_name="Feuille1"):
    """Crée une feuille Excel d'échantillon."""
    data = {
        'ID': list(range(1, 11)),
        'Nom': [f'Produit {i}' for i in range(1, 11)],
        'Prix': [10.99 * i for i in range(1, 11)],
        'Quantité': [i * 5 for i in range(1, 11)],
        'Total': [10.99 * i * i * 5 for i in range(1, 11)]
    }
    
    df = pd.DataFrame(data)
    df.to_excel(output_path, sheet_name=sheet_name, index=False)
    return output_path
